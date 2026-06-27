"""
export_docx.py — paper.md を Word(.docx) に変換するスクリプト
画像を埋め込み、単体ファイルで完結させる。

実行方法：
    cd selfish_genes
    python src/export_docx.py
出力：docs/paper.docx
"""
import os, re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR   = os.path.dirname(os.path.dirname(__file__))
MD_PATH    = os.path.join(BASE_DIR, "docs", "paper.md")
OUT_PATH   = os.path.join(BASE_DIR, "docs", "paper.docx")
FIGS_DIR   = os.path.join(BASE_DIR, "data", "figures")

def add_heading(doc, text, level):
    doc.add_heading(text, level=level)

def add_paragraph(doc, text):
    """太字・イタリック記法を処理してパラグラフを追加"""
    p = doc.add_paragraph()
    # **bold** と *italic* を交互に処理
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            # $数式$ はそのままテキストとして挿入
            p.add_run(part)
    return p

def add_table_from_md(doc, lines):
    """Markdown テーブルを docx テーブルに変換"""
    rows = []
    for line in lines:
        if re.match(r'\|[-| :]+\|', line):
            continue  # 区切り行スキップ
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = 'Table Grid'
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            if j < ncols:
                tc = table.rows[i].cells[j]
                tc.text = cell
                if i == 0:
                    for run in tc.paragraphs[0].runs:
                        run.bold = True

def main():
    doc = Document()

    # 余白設定
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    with open(MD_PATH, encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')

        # --- 見出し ---
        if line.startswith('#### '):
            add_heading(doc, line[5:], 4); i += 1; continue
        if line.startswith('### '):
            add_heading(doc, line[4:], 3); i += 1; continue
        if line.startswith('## '):
            add_heading(doc, line[3:], 2); i += 1; continue
        if line.startswith('# '):
            add_heading(doc, line[2:], 1); i += 1; continue

        # --- 水平線 ---
        if re.match(r'^---+$', line):
            i += 1; continue

        # --- 画像 ![...](path) ---
        img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line)
        if img_match:
            alt, rel_path = img_match.group(1), img_match.group(2)
            # ../data/figures/figN.png → data/figures/figN.png
            img_path = os.path.normpath(os.path.join(BASE_DIR, "docs", rel_path))
            if os.path.exists(img_path):
                doc.add_picture(img_path, width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_paragraph(f"[画像なし: {rel_path}]")
            i += 1; continue

        # --- テーブル ---
        if line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i].rstrip('\n'))
                i += 1
            add_table_from_md(doc, table_lines)
            doc.add_paragraph()
            continue

        # --- 参考文献（箇条書き "- "）---
        if line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(re.sub(r'[*_`]', '', line[2:]))
            i += 1; continue

        # --- イタリックキャプション行（*図N:...* ）---
        if line.startswith('*') and line.endswith('*') and len(line) > 2:
            p = doc.add_paragraph()
            run = p.add_run(line[1:-1])
            run.italic = True
            run.font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1; continue

        # --- 空行 ---
        if line.strip() == '':
            i += 1; continue

        # --- 通常段落 ---
        add_paragraph(doc, line)
        i += 1

    doc.save(OUT_PATH)
    print(f"保存完了: {OUT_PATH}")

if __name__ == "__main__":
    main()
